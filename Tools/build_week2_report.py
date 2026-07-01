# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Week2_Report_Modeng.docx").resolve()


def rfont(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def ptext(paragraph, text, size=11, bold=False, color=None):
    run = paragraph.add_run(text)
    rfont(run, size, bold, color)
    return run


def cell_text(cell, text, size=10.5, bold=False, color=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    if align:
        paragraph.alignment = align
    ptext(paragraph, text, size, bold, color)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def borders(table, color="D0D7DE"):
    props = table._tbl.tblPr
    border_node = props.find(qn("w:tblBorders"))
    if border_node is None:
        border_node = OxmlElement("w:tblBorders")
        props.append(border_node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = border_node.find(qn("w:" + edge))
        if item is None:
            item = OxmlElement("w:" + edge)
            border_node.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    props = table._tbl.tblPr
    width_node = props.find(qn("w:tblW"))
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        props.append(width_node)
    width_node.set(qn("w:w"), str(sum(widths)))
    width_node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            props = cell._tc.get_or_add_tcPr()
            cell_width = props.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                props.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[idx]))
            cell_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def margins(table, top=80, start=120, bottom=80, end=120):
    props = table._tbl.tblPr
    margin_node = props.find(qn("w:tblCellMar"))
    if margin_node is None:
        margin_node = OxmlElement("w:tblCellMar")
        props.append(margin_node)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = margin_node.find(qn("w:" + key))
        if item is None:
            item = OxmlElement("w:" + key)
            margin_node.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    ptext(paragraph, text, 16 if level == 1 else 13, True, "2E74B5" if level <= 2 else "1F4D78")


def body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    ptext(paragraph, text, 11)


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    ptext(paragraph, text, 11)


def build():
    doc = Document()
    section = doc.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "《墨灯守巷》第二周工作总结报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        rfont(run, 9, False, "666666")

    footer = section.footer.paragraphs[0]
    footer.text = "第1组 | Unreal Engine 5.7 项目周报"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        rfont(run, 9, False, "666666")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    ptext(title, "每周总结报告", 22, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    ptext(subtitle, "《墨灯守巷》第二周项目进度总结", 12, False, "555555")

    meta = doc.add_table(3, 2)
    borders(meta)
    margins(meta)
    table_width(meta, [1800, 7560])
    for idx, (key, value) in enumerate((("组别", "第1组"), ("游戏名", "墨灯守巷"), ("时间周期", "第二周"))):
        shade(meta.cell(idx, 0), "F2F4F7")
        cell_text(meta.cell(idx, 0), key, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(meta.cell(idx, 1), value)

    doc.add_paragraph()
    heading(doc, "模块状态", 1)

    modules = [
        ("进行中", "黄帝尧", "怪物逻辑与交互系统", "根据项目计划书第二周目标，本周围绕“守灯笼、打敌人、胜负判定”的最小玩法闭环展开开发。主要完成怪物基础逻辑、敌人血条、灯笼交互提示，以及敌人动画接入的 C++ 预留；当前项目已具备敌人攻击灯笼、玩家修复灯笼、波次刷怪、胜利/失败判断等核心规则，并通过 GitHub 与 Git LFS 进行同步管理。"),
        ("进行中", "邢增浩", "主角移动、跳跃与攻击", "本周主要负责主角移动、跳跃和攻击相关开发。玩家已具备横板移动、跳跃、普通攻击和基础攻击范围检测能力，可通过鼠标左键或 J 攻击敌人；相关输入与角色动作逻辑为后续接入 Tomoe 主角模型、剑术攻击动画和更完整的战斗反馈打下基础。"),
        ("已结束", "黄顺金", "第一关原型与敌人视觉资源", "本周已搭建第一关“灯会街口”原型场景，并接入 Stickman 敌人资源。基础敌人已从占位模型升级为 Stickman 骨骼模型，并具备待机、行走、攻击、受击、死亡等基础动画表现。快速敌人与自爆敌人蓝图也已进行阶段性更新，但仍保留进一步美术区分空间。"),
        ("进行中", "张乐钊", "UI、提示与测试支持", "本周已完成 C++ HUD 阶段性显示，包括当前波次、敌人数量、灯笼状态、武器等级、墨点、伤害和攻击范围等信息；新增敌人血条和灯笼修复提示，玩家靠近受损灯笼时会显示 E / F Repair。胜利和失败结果界面已具备重开与退出功能，后续仍需向正式 UMG 界面、美术化 UI 和音效提示推进。"),
        ("进行中", "李远涵", "技术展示与后续渲染准备", "本周主要配合项目核心闭环完成，暂未进入自定义后处理正式实现阶段。当前灯笼系统已经具备耐久、明暗和颜色变化逻辑，为后续实现“灯笼越暗画面越冷灰、灯笼恢复画面变暖变亮”的后处理联动效果提供了玩法数据基础。"),
    ]

    module_table = doc.add_table(1, 4)
    borders(module_table)
    margins(module_table, 100, 120, 100, 120)
    table_width(module_table, [1100, 1100, 1700, 5460])
    for idx, header_text in enumerate(("模块状态", "负责人", "模块", "详细说明")):
        shade(module_table.cell(0, idx), "E8EEF5")
        cell_text(module_table.cell(0, idx), header_text, bold=True, color="1F3A5F", align=WD_ALIGN_PARAGRAPH.CENTER)

    for status, owner, module, detail in modules:
        row = module_table.add_row().cells
        cell_text(row[0], status, bold=True, color="1F3A5F", align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[1], owner, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[2], module, size=10, bold=True)
        cell_text(row[3], detail, size=9.5)

    heading(doc, "本周小结", 1)
    heading(doc, "小组总结", 2)
    for text in (
        "本周开发重点围绕项目计划书中第二周阶段验收标准展开，即完成第一关灰盒原型、玩家攻击、敌人生成、敌人受击与死亡、灯笼耐久系统以及基础胜负判断。当前项目已经形成较完整的第一关最小玩法闭环：敌人会按波次生成并向灯笼移动，攻击灯笼会降低耐久；玩家可以移动、跳跃、攻击敌人、修复灯笼；击败敌人可以获得墨点并提升武器能力；所有灯笼熄灭会失败，清除波次可胜利。",
        "在表现层面，本周完成了 Stickman 敌人资源导入与基础动画接入，基础敌人不再只使用占位模型；同时开始导入 Tomoe 主角模型和剑术动画资源，为后续主角替换和战斗表现优化打基础。UI 方面已经具备调试型 HUD、敌人血条、灯笼修复提示和胜负结果界面。版本管理方面，团队已使用 GitHub 与 Git LFS 管理 Unreal 工程和大体积美术资源，并在开发中明确了避免多人同时修改同一 .uasset / .umap 文件的协作规则。",
        "总体来看，第二周目标基本完成，并且部分内容已经超出原计划进入第三周的波次、UI、动画表现和资源替换阶段。当前不足主要集中在：第一关仍偏灰盒与原型状态，UI 仍以 C++ Canvas HUD 为主，音效、粒子、玩家血量、灯火弹、正式暂停菜单和后处理效果尚未完善；快速敌人和自爆敌人的独立美术表现仍需继续区分；主角 Tomoe 的移动动画和攻击动画还需要继续调试。",
    ):
        body(doc, text)

    heading(doc, "组员个人小结", 2)
    member_table = doc.add_table(6, 2)
    borders(member_table)
    margins(member_table)
    table_width(member_table, [1600, 7760])
    for idx, header_text in enumerate(("组员", "小结")):
        shade(member_table.cell(0, idx), "E8EEF5")
        cell_text(member_table.cell(0, idx), header_text, bold=True, color="1F3A5F", align=WD_ALIGN_PARAGRAPH.CENTER)
    member_summaries = (
        ("黄帝尧", "本周主要负责怪物逻辑、敌人血条、灯笼交互提示和敌人动画接入的 C++ 预留。围绕第二周“最小玩法闭环”目标，推进敌人生成、移动、攻击灯笼、受击死亡、血条显示、修复提示和动画接口预留等内容，确保第一关原型能够形成可运行的守护玩法。"),
        ("邢增浩", "本周主要负责主角移动、跳跃和攻击的开发。完成并完善横板角色移动、跳跃输入、普通攻击触发和攻击检测等基础能力，使玩家可以在第一关原型中完成移动、躲避和攻击敌人的核心操作。"),
        ("黄顺金", "本周主要负责关卡原型和美术资源接入。完成第一关灰盒场景的基础搭建，整理并导入 Stickman 敌人资源，调整基础敌人外观和动画表现，使敌人从占位模型过渡到可展示的骨骼模型效果。"),
        ("张乐钊", "本周主要关注 UI、提示、测试和运行反馈。当前 HUD 已能显示波次、敌人、灯笼、武器和墨点等关键状态，胜负结果界面、敌人血条和灯笼修复提示也已能辅助玩家理解当前战斗情况，后续将继续向正式 UMG UI 和音效提示完善。"),
        ("李远涵", "本周主要配合核心玩法闭环与后续技术展示准备。虽然自定义后处理尚未正式实现，但灯笼耐久、明暗变化和颜色反馈已经具备基础数据来源，后续可以在此基础上继续实现灯笼状态与画面冷暖、饱和度和暗角变化的联动效果。"),
    )
    for idx, (name, summary) in enumerate(member_summaries, 1):
        cell_text(member_table.cell(idx, 0), name, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(member_table.cell(idx, 1), summary, size=9.8)

    heading(doc, "小组讨论纪要和必要截图", 1)
    for text in (
        "本周讨论确定继续以“最小玩法闭环”为优先目标，先保证玩家、敌人、灯笼、波次和胜负判断可以完整运行，再逐步替换美术、动画、UI、音效和后处理效果。",
        "团队明确了 C++ 与蓝图/资源分工：C++ 侧主要负责玩家攻击、敌人逻辑、灯笼耐久、波次刷怪、胜负判定、HUD 和交互提示；蓝图和资源侧主要负责敌人外观、主角模型、动画资源、关卡摆放和视觉表现。",
        "为避免 Git LFS 冲突，讨论并确定协作规则：同一时间不要两个人同时修改同一个 .uasset 或 .umap 文件。尤其是 BP_ModengEnemy、BP_ModengFastEnemy、BP_ModengExploderEnemy、L_Level01_Street 以及 Content/__ExternalActors__ 下的关卡外部 Actor 文件，修改前需要先在群里说明“锁定”对应资源，提交推送后其他成员再拉取继续工作。",
        "本周 GitHub 仓库已同步多次更新，主要提交包括敌人 Stickman 资源接入、基础敌人动画改进、敌人血条与灯笼修复提示、项目 README 与交接文档更新、Tomoe 主角资源和战斗动画资源接入等。",
    ):
        bullet(doc, text)

    heading(doc, "截图建议", 2)
    shot_table = doc.add_table(1, 2)
    borders(shot_table)
    margins(shot_table, 100, 120, 100, 120)
    table_width(shot_table, [1500, 7860])
    for idx, header_text in enumerate(("序号", "截图内容")):
        shade(shot_table.cell(0, idx), "F2F4F7")
        cell_text(shot_table.cell(0, idx), header_text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, text in enumerate((
        "第一关灰盒场景运行截图。",
        "玩家攻击敌人截图。",
        "敌人血条显示截图。",
        "灯笼受损后显示 E / F Repair 提示截图。",
        "胜利或失败结果界面截图。",
        "Stickman 敌人动画或 Tomoe 主角模型接入截图。",
    ), 1):
        row = shot_table.add_row().cells
        cell_text(row[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[1], text)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
