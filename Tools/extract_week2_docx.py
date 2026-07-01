from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            para = None
            for candidate in document.paragraphs:
                if candidate._p is child:
                    para = candidate
                    break
            if para is not None:
                text = para.text.strip()
                if text:
                    yield {"type": "paragraph", "text": text}
        elif child.tag.endswith("}tbl"):
            table = None
            for candidate in document.tables:
                if candidate._tbl is child:
                    table = candidate
                    break
            if table is not None:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                yield {"type": "table", "rows": rows}


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: extract_week2_docx.py INPUT.docx OUTPUT.json", file=sys.stderr)
        return 2

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    doc = Document(str(input_path))
    blocks = list(iter_block_items(doc))
    payload = {
        "source": str(input_path),
        "paragraphs": [b["text"] for b in blocks if b["type"] == "paragraph"],
        "tables": [b["rows"] for b in blocks if b["type"] == "table"],
        "blocks": blocks,
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"wrote {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
